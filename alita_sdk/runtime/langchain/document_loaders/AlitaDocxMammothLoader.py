import re
import uuid
from io import BytesIO

import mammoth.images
import pytesseract
from PIL import Image
from langchain_core.document_loaders import BaseLoader
from langchain_core.documents import Document
from mammoth import convert_to_html
from markdownify import markdownify
from docx import Document as DocxDocument
from docx.oxml.ns import qn
from bs4 import BeautifulSoup

from alita_sdk.tools.chunkers.sematic.markdown_chunker import markdown_by_headers_chunker
from .utils import perform_llm_prediction_for_image_bytes


class AlitaDocxMammothLoader(BaseLoader):
    """
    Loader for Docx files using Mammoth to convert to HTML, with image handling,
    and then Markdownify to convert HTML to markdown.
    Detects bordered paragraphs and text boxes and treats them as code blocks.
    """
    def __init__(self, **kwargs):
        """
        Initializes AlitaDocxMammothLoader.

        Args:
            **kwargs: Keyword arguments, including:
                file_path (str): Path to the Docx file. Required.
                llm (LLM, optional): Language model for processing images.
                prompt (str, optional): Prompt for the language model.
        Raises:
            ValueError: If the 'path' parameter is not provided.
        """
        self.path =  kwargs.get('file_path')
        self.file_content = kwargs.get('file_content')
        self.file_name = kwargs.get('file_name')
        self.extract_images = kwargs.get('extract_images')
        self.llm = kwargs.get("llm")
        self.prompt = kwargs.get("prompt")
        self.max_tokens = kwargs.get('max_tokens', 512)

    def __handle_image(self, image) -> dict:
        """
        Handles image processing within the Docx file.

        Uses LLM for image captioning if available, otherwise falls back to Tesseract OCR,
        and defaults to base64 encoding if both fail.

        Args:
            image (mammoth.images.Image): Image object from Mammoth.

        Returns:
            dict: Dictionary containing the 'src' attribute for the HTML <img> tag.
        """
        output = {}
        try:
            if self.llm:
                # Use LLM for image understanding
                with image.open() as image_bytes:
                    result = perform_llm_prediction_for_image_bytes(image_bytes, self.llm, self.prompt)
                output['src'] = result  # LLM image transcript in src
                return output
            else:
                # Use Tesseract OCR if LLM is not available
                with image.open() as image_bytes:
                    img = Image.open(image_bytes)
                    output['src'] = pytesseract.image_to_string(image=img)  # Tesseract transcript in src
                return output
        except Exception as e:
            # Fallback to default image handling for any exceptions during image processing
            # This ensures robustness against various image format issues, OCR failures,
            # or LLM invocation problems. It prevents the loader from crashing due to
            # a single image processing failure and provides a default image representation.
            return self.__default_image_handler(image)

    def __default_image_handler(self, image) -> dict:
        """
        Default image handler: encodes image to base64 data URL.

        Args:
            image (mammoth.images.Image): Image object from Mammoth.

        Returns:
            dict: Dictionary with base64 encoded 'src' for HTML <img> tag.
        """
        return {"src": "Transcript is not available"}


    def __postprocess_original_md(self, original_md: str) -> str:
        # Pattern to match placeholders like[image_1_1.png] or similar
        pattern = re.compile(r'!\[([^\]]*)\]\(([^)]*)(\s\"([^"]*)\")?\)')

        def replace_placeholder(match):
            transcript = match.group(2)
            # Return a markdown formatted transcript section.
            return f"\n**Image Transcript:**\n{transcript}\n"

        new_md = pattern.sub(replace_placeholder, original_md)
        return new_md

    def __has_border(self, paragraph):
        """
        Check if a paragraph has border formatting.
        
        Args:
            paragraph: A python-docx Paragraph object.
            
        Returns:
            bool: True if paragraph has any border, False otherwise.
        """
        pPr = paragraph._element.pPr
        if pPr is not None:
            pBdr = pPr.find(qn('w:pBdr'))
            if pBdr is not None:
                # Check if any border side exists (top, bottom, left, right)
                for side in ['top', 'bottom', 'left', 'right']:
                    border = pBdr.find(qn(f'w:{side}'))
                    if border is not None:
                        # Check if border is not "none" or has a width
                        val = border.get(qn('w:val'))
                        if val and val != 'none':
                            return True
        return False

    def __find_text_boxes(self, doc):
        """
        Find all text boxes in document by searching OOXML structure.
        Text boxes are typically in w:txbxContent elements.
        
        Args:
            doc: A python-docx Document object.
            
        Returns:
            list: List of tuples (element, paragraphs_inside_textbox).
        """
        text_boxes = []
        
        # Iterate through document body XML to find text box content elements
        for element in doc.element.body.iter():
            # Look for text box content elements
            if element.tag.endswith('txbxContent'):
                # Collect all paragraphs inside this text box
                txbx_paragraphs = []
                for txbx_para_element in element.iter():
                    if txbx_para_element.tag.endswith('p'):
                        txbx_paragraphs.append(txbx_para_element)
                
                if txbx_paragraphs:
                    text_boxes.append((element, txbx_paragraphs))
        
        return text_boxes

    def __create_marker_paragraph(self, marker_text):
        """
        Create a paragraph element with marker text.
        
        Args:
            marker_text (str): The marker text to insert.
            
        Returns:
            Element: An OOXML paragraph element.
        """
        from docx.oxml import OxmlElement
        
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = marker_text
        r.append(t)
        p.append(r)
        return p

    def __inject_markers_for_paragraph(self, paragraph, start_marker, end_marker):
        """
        Inject marker paragraphs before and after a bordered paragraph.
        
        Args:
            paragraph: A python-docx Paragraph object.
            start_marker (str): The start marker text.
            end_marker (str): The end marker text.
        """
        # Insert start marker paragraph before
        marker_p_start = self.__create_marker_paragraph(start_marker)
        paragraph._element.addprevious(marker_p_start)
        
        # Insert end marker paragraph after
        marker_p_end = self.__create_marker_paragraph(end_marker)
        paragraph._element.addnext(marker_p_end)

    def __inject_markers_for_textbox(self, textbox_element, paragraph_elements, start_marker, end_marker):
        """
        Inject markers around text box content.
        
        Args:
            textbox_element: The w:txbxContent element.
            paragraph_elements: List of paragraph elements inside the text box.
            start_marker (str): The start marker text.
            end_marker (str): The end marker text.
        """
        if not paragraph_elements:
            return
        
        # Insert start marker before first paragraph in text box
        first_para = paragraph_elements[0]
        marker_p_start = self.__create_marker_paragraph(start_marker)
        first_para.addprevious(marker_p_start)
        
        # Insert end marker after last paragraph in text box
        last_para = paragraph_elements[-1]
        marker_p_end = self.__create_marker_paragraph(end_marker)
        last_para.addnext(marker_p_end)

    def __detect_and_mark_bordered_content(self, docx_stream):
        """
        Detects bordered paragraphs and text boxes, injects unique markers around them.
        Groups consecutive bordered paragraphs into single code blocks.
        
        Args:
            docx_stream: A file-like object containing the DOCX document.
            
        Returns:
            tuple: (modified_docx_stream, start_marker, end_marker)
        """
        # Load document with python-docx
        doc = DocxDocument(docx_stream)
        
        # Generate unique markers to avoid conflicts with document content
        unique_id = uuid.uuid4().hex[:8]
        start_marker = f"<<<BORDERED_BLOCK_START_{unique_id}>>>"
        end_marker = f"<<<BORDERED_BLOCK_END_{unique_id}>>>"
        
        # Group consecutive bordered paragraphs together
        bordered_groups = []
        current_group = []
        
        for para in doc.paragraphs:
            if self.__has_border(para):
                current_group.append(para)
            else:
                if current_group:
                    # End of a bordered group
                    bordered_groups.append(current_group)
                    current_group = []
        
        # Don't forget the last group if document ends with bordered paragraphs
        if current_group:
            bordered_groups.append(current_group)
        
        # Collect all text boxes
        # text_boxes = self.__find_text_boxes(doc)
        
        # Inject markers around each group of consecutive bordered paragraphs
        for group in bordered_groups:
            if group:
                # Add start marker before first paragraph in group
                first_para = group[0]
                marker_p_start = self.__create_marker_paragraph(start_marker)
                first_para._element.addprevious(marker_p_start)
                
                # Add end marker after last paragraph in group
                last_para = group[-1]
                marker_p_end = self.__create_marker_paragraph(end_marker)
                last_para._element.addnext(marker_p_end)
        
        # Inject markers around text box content
        # for textbox_element, para_elements in text_boxes:
        #     self.__inject_markers_for_textbox(textbox_element, para_elements, start_marker, end_marker)
        
        # Save modified document to BytesIO
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        
        return output, start_marker, end_marker

    def __contains_complex_structure(self, content_html):
        """
        Check if HTML content contains tables, lists, or other complex structures.
        
        Args:
            content_html (str): HTML content to analyze.
            
        Returns:
            bool: True if content contains tables/lists, False otherwise.
        """
        content_soup = BeautifulSoup(content_html, 'html.parser')
        
        # Check for tables
        if content_soup.find('table'):
            return True
        
        # Check for lists (ul, ol)
        if content_soup.find('ul') or content_soup.find('ol'):
            return True
        
        return False

    def __escape_hash_symbols(self, html_content):
        """
        Escape hash (#) symbols at the beginning of lines in HTML to prevent
        them from being treated as markdown headers.
        
        Args:
            html_content (str): HTML content.
            
        Returns:
            str: HTML with escaped hash symbols.
        """
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Process all text-containing elements
        for element in soup.find_all(['p', 'li', 'td', 'th', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
            if element.string:
                text = element.string
                # If line starts with #, escape it
                if text.strip().startswith('#'):
                    element.string = text.replace('#', '\\#', 1)
        
        return str(soup)

    def __wrap_marked_sections_in_code_blocks(self, html, start_marker, end_marker):
        """
        Find content between markers and wrap appropriately:
        - Simple text/code → <pre><code> block
        - Tables/lists → Custom wrapper with preserved structure
        
        Args:
            html (str): The HTML content from Mammoth.
            start_marker (str): The start marker text.
            end_marker (str): The end marker text.
            
        Returns:
            str: HTML with marked sections wrapped appropriately.
        """
        import html as html_module
        
        # Mammoth escapes < and > to &lt; and &gt;, so we need to escape our markers too
        escaped_start = html_module.escape(start_marker)
        escaped_end = html_module.escape(end_marker)
        
        # Pattern to find content between HTML-escaped markers (including HTML tags)
        # The markers will be in separate <p> tags, and content in between
        pattern = re.compile(
            f'<p>{re.escape(escaped_start)}</p>(.*?)<p>{re.escape(escaped_end)}</p>',
            re.DOTALL
        )
        
        def replace_with_appropriate_wrapper(match):
            content = match.group(1)
            
            # Detect if content has complex structure (tables, lists)
            has_complex_structure = self.__contains_complex_structure(content)
            
            if has_complex_structure:
                # Preserve structure: keep HTML as-is, escape # symbols
                escaped_content = self.__escape_hash_symbols(content)
                # Wrap in a div with special class for potential custom handling
                return f'<div class="alita-bordered-content">{escaped_content}</div>'
            else:
                # Simple text/code: extract as plain text and wrap in code block
                content_soup = BeautifulSoup(content, 'html.parser')
                
                # Extract text from each paragraph separately to preserve line breaks
                lines = []
                for element in content_soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6']):
                    # Replace <br /> within paragraphs with newlines
                    for br in element.find_all('br'):
                        br.replace_with('\n')
                    text = element.get_text()
                    # Preserve leading whitespace (indentation), only strip trailing
                    lines.append(text.rstrip())
                
                # If no paragraphs found, just get all text
                if not lines:
                    content = content.replace('<br />', '\n').replace('<br/>', '\n').replace('<br>', '\n')
                    content_text = content_soup.get_text()
                    lines = [line.rstrip() for line in content_text.split('\n')]
                
                # Join lines, strip only leading/trailing empty lines
                content_text = '\n'.join(lines).strip()
                # Return as code block (need to HTML-escape the content)
                content_escaped = html_module.escape(content_text)
                return f'<pre><code>{content_escaped}</code></pre>'
        
        # Replace all marked sections with appropriate wrappers
        result_html = pattern.sub(replace_with_appropriate_wrapper, html)
        
        return result_html

    def load(self):
        """
        Loads and converts the Docx file to markdown format.

        Returns:
            List[Document]: A list containing a Documents with the markdown content
                          and metadata including the source file path.
        """
        result_content = self.get_content()
        return list(markdown_by_headers_chunker(iter([Document(page_content=result_content, metadata={'source': str(self.path)})]), config={'max_tokens':self.max_tokens}))

    def get_content(self):
        """
        Extracts and converts the content of the Docx file to markdown format.

        Handles both file paths and in-memory file content.

        Returns:
            str: The markdown content extracted from the Docx file.
        """
        if self.path:
            # If path is provided, read from file system
            with open(self.path, 'rb') as docx_file:
                return self._convert_docx_to_markdown(docx_file)
        elif self.file_content and self.file_name:
            # If file_content and file_name are provided, read from memory
            docx_file = BytesIO(self.file_content)
            return self._convert_docx_to_markdown(docx_file)
        else:
            raise ValueError("Either 'path' or 'file_content' and 'file_name' must be provided.")

    def _convert_docx_to_markdown(self, docx_file):
        """
        Converts the content of a Docx file to markdown format.
        Detects bordered content and treats it as code blocks.

        Args:
            docx_file (BinaryIO): The Docx file object.

        Returns:
            str: The markdown content extracted from the Docx file.
        """
        # Step 1: Detect and mark bordered content
        # Reset stream position if needed
        if hasattr(docx_file, 'seek'):
            docx_file.seek(0)
        
        marked_docx, start_marker, end_marker = self.__detect_and_mark_bordered_content(docx_file)
        
        # Step 2: Convert marked DOCX to HTML using Mammoth
        if self.extract_images:
            # Extract images using the provided image handler
            result = convert_to_html(marked_docx, convert_image=mammoth.images.img_element(self.__handle_image))
        else:
            # Ignore images
            result = convert_to_html(marked_docx, convert_image=lambda image: "")
        
        # Step 3: Wrap marked sections in <pre><code> tags
        html_with_code_blocks = self.__wrap_marked_sections_in_code_blocks(
            result.value, start_marker, end_marker
        )
        
        # Step 4: Convert HTML to markdown
        content = markdownify(html_with_code_blocks, heading_style="ATX")
        
        # Step 5: Post-process markdown (for image transcripts, etc.)
        return self.__postprocess_original_md(content)
